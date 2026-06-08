"""Document service for handling uploads, chunking, and embeddings."""
import os
from pathlib import Path
from uuid import UUID
from typing import List, Dict, Any
import PyPDF2
import weaviate

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.models.db import Document
from app.config import settings
from app.utils.logging import get_logger
from app.utils.errors import DocumentNotFoundError, VectorDatabaseError

logger = get_logger(__name__)


class DocumentChunker:
    """Split documents into chunks for embedding."""

    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        chunks = []
        words = text.split()

        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk = " ".join(words[i : i + self.chunk_size])
            if chunk.strip():
                chunks.append(chunk)

        return chunks


class EmbeddingGenerator:
    """Generate embeddings using sentence-transformers."""

    def __init__(self):
        try:
            from sentence_transformers import SentenceTransformer

            self.model = SentenceTransformer(settings.EMBEDDING_MODEL)
        except ImportError:
            logger.error("sentence-transformers not installed")
            raise

    def embed(self, text: str) -> List[float]:
        """Generate embedding for text."""
        try:
            embedding = self.model.encode(text, convert_to_tensor=False)
            return embedding.tolist()
        except Exception as e:
            logger.error(f"Error generating embedding: {str(e)}")
            raise VectorDatabaseError(f"Embedding generation failed: {str(e)}")


class DocumentService:
    """Service for document management."""

    def __init__(self):
        self.chunker = DocumentChunker(chunk_size=500, overlap=50)
        self.embedder = EmbeddingGenerator()
        self.weaviate_client = weaviate.Client(settings.WEAVIATE_URL)
        self.storage_path = Path(settings.STORAGE_PATH)
        self.storage_path.mkdir(parents=True, exist_ok=True)

    async def upload_document(
        self,
        file_content: bytes,
        filename: str,
        tenant_id: UUID,
        user_id: UUID,
        db: AsyncSession,
    ) -> Document:
        """
        Upload a document and store metadata.
        Actual chunking and embedding happens in async task.
        """
        try:
            # Create storage path
            tenant_dir = self.storage_path / str(tenant_id)
            tenant_dir.mkdir(parents=True, exist_ok=True)

            # Save file
            file_path = tenant_dir / filename
            with open(file_path, "wb") as f:
                f.write(file_content)

            # Create database record
            storage_path = str(file_path.relative_to(self.storage_path))
            document = Document(
                tenant_id=tenant_id,
                original_filename=filename,
                storage_path=storage_path,
                document_size_bytes=len(file_content),
                uploaded_by=user_id,
                processed=False,  # Will be set to True after async processing
                metadata={"file_type": self._get_file_type(filename)},
            )

            db.add(document)
            await db.commit()
            await db.refresh(document)

            logger.info(
                f"Uploaded document {document.id} for tenant {tenant_id}: {filename}"
            )

            return document

        except Exception as e:
            logger.error(f"Error uploading document: {str(e)}")
            await db.rollback()
            raise

    async def process_document(
        self, document_id: UUID, tenant_id: UUID, db: AsyncSession
    ) -> None:
        """
        Process document: extract text, chunk, embed, and store in Weaviate.
        This should be called as an async task after upload.
        """
        try:
            # Get document
            result = await db.execute(
                select(Document).where(Document.id == document_id)
            )
            document = result.scalar_one_or_none()

            if not document:
                raise DocumentNotFoundError(f"Document {document_id} not found")

            # Extract text
            file_path = self.storage_path / document.storage_path
            text = self._extract_text(file_path)

            if not text.strip():
                logger.warning(f"No text extracted from document {document_id}")
                document.processed = True
                await db.commit()
                return

            # Chunk text
            chunks = self.chunker.chunk_text(text)

            # Get tenant's Weaviate collection
            result = await db.execute(
                select(Document).where(Document.id == document_id)
            )
            doc = result.scalar_one_or_none()
            if not doc or not doc.tenant:
                raise DocumentNotFoundError("Document or tenant not found")

            # Store chunks in Weaviate
            self._store_chunks_in_weaviate(
                chunks=chunks,
                document_id=str(document_id),
                document_name=document.original_filename,
                tenant_id=str(tenant_id),
                collection_name=doc.tenant.vector_db_collection_name,
            )

            # Mark as processed
            document.processed = True
            await db.commit()

            logger.info(f"Processed document {document_id}: {len(chunks)} chunks")

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            await db.rollback()
            raise

    def _extract_text(self, file_path: Path) -> str:
        """Extract text from file (PDF, TXT, DOCX)."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_path.suffix.lower()

        try:
            if file_type == ".pdf":
                return self._extract_pdf(file_path)
            elif file_type == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            elif file_type == ".docx":
                return self._extract_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF."""
        try:
            text = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            raise

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            return "\n".join(text)
        except ImportError:
            logger.warning("python-docx not installed, cannot extract DOCX")
            return ""
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            raise

    def _store_chunks_in_weaviate(
        self,
        chunks: List[str],
        document_id: str,
        document_name: str,
        tenant_id: str,
        collection_name: str,
    ) -> None:
        """Store document chunks with embeddings in Weaviate."""
        try:
            for i, chunk in enumerate(chunks):
                # Generate embedding
                embedding = self.embedder.embed(chunk)

                # Create object
                data_object = {
                    "content": chunk,
                    "source": document_name,
                    "document_id": document_id,
                    "chunk_index": i,
                    "tenant_id": tenant_id,
                }

                # Add to Weaviate with vector
                self.weaviate_client.data_object.create(
                    data_object=data_object,
                    class_name=collection_name,
                    vector=embedding,
                )

            logger.info(
                f"Stored {len(chunks)} chunks in Weaviate for document {document_id}"
            )

        except Exception as e:
            logger.error(f"Error storing chunks in Weaviate: {str(e)}")
            raise VectorDatabaseError(f"Failed to store chunks: {str(e)}")

    async def get_document(
        self, document_id: UUID, db: AsyncSession
    ) -> Document:
        """Get document by ID."""
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            raise DocumentNotFoundError(f"Document {document_id} not found")

        return document

    async def list_documents(
        self, tenant_id: UUID, db: AsyncSession
    ) -> List[Document]:
        """List all documents for a tenant."""
        result = await db.execute(
            select(Document).where(Document.tenant_id == tenant_id)
        )
        return result.scalars().all()

    async def delete_document(
        self, document_id: UUID, tenant_id: UUID, db: AsyncSession
    ) -> None:
        """Delete a document and its embeddings."""
        try:
            document = await self.get_document(document_id, db)

            # Verify ownership
            if document.tenant_id != tenant_id:
                raise ValueError("Document does not belong to this tenant")

            # Delete from Weaviate
            try:
                where_filter = {
                    "path": ["document_id"],
                    "operator": "Equal",
                    "valueString": str(document_id),
                }
                self.weaviate_client.data_object.delete(
                    where=where_filter, class_name=document.tenant.vector_db_collection_name
                )
            except Exception as e:
                logger.warning(f"Error deleting from Weaviate: {str(e)}")

            # Delete file
            file_path = self.storage_path / document.storage_path
            if file_path.exists():
                file_path.unlink()

            # Delete from database
            await db.delete(document)
            await db.commit()

            logger.info(f"Deleted document {document_id}")

        except Exception as e:
            logger.error(f"Error deleting document: {str(e)}")
            await db.rollback()
            raise

    @staticmethod
    def _get_file_type(filename: str) -> str:
        """Get file type from filename."""
        return filename.split(".")[-1].lower() if "." in filename else "unknown"


# Create singleton instance
document_service = DocumentService()
